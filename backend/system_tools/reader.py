# backend/system_tools/reader.py
import base64
import asyncio
import mimetypes
from pathlib import Path
from typing import Optional, Union, Dict, Any, List
from dataclasses import dataclass, asdict
from functools import lru_cache

from config_loader import config
from backend.utils.validators import validate_path
import backend


# ============ 自定义异常 ============
class FileReadError(Exception):
    """
    文件读取过程中出现的错误，将被 MCP 框架转为 isError 响应。
    """
    def __init__(self, message: str, code: str = 'READ_FAILED'):
        super().__init__(message)
        self.code = code


# ============ 数据结构 ============
@dataclass
class FileReadResult:
    """文件读取结果"""
    content: str
    format: str          # 原始格式: text, markdown, json, csv, binary, image 等
    mime_type: str
    metadata: Dict[str, Any]  # 额外元信息（如 sheet 名、页数、图片尺寸等）


@dataclass
class ReadContext:
    """传递给各个处理器的上下文对象"""
    path: Path
    mime_type: str
    encoding: Optional[str] = None
    sheet_name: Optional[Union[str, int]] = None
    return_base64: bool = False
    file_size: int = 0


# ============ 可选依赖的懒加载与缓存 ============
@lru_cache(maxsize=1)
def _get_markitdown():
    try:
        from markitdown import MarkItDown
        return MarkItDown()
    except ImportError:
        return None

@lru_cache(maxsize=1)
def _get_pandas():
    try:
        import pandas as pd
        return pd
    except ImportError:
        return None

@lru_cache(maxsize=1)
def _get_docx():
    try:
        import docx
        return docx
    except ImportError:
        return None

@lru_cache(maxsize=1)
def _get_pypdf():
    try:
        import PyPDF2
        return PyPDF2
    except ImportError:
        return None

@lru_cache(maxsize=1)
def _get_chardet():
    try:
        import chardet
        return chardet
    except ImportError:
        return None


# ============ 通用工具函数 ============
def df_to_markdown(df, max_rows: int = 2000) -> str:
    """将 DataFrame 转为 Markdown 表格，使用 numpy 加速并限制最大行数"""
    n_rows = len(df)
    truncated = n_rows > max_rows
    df_show = df.head(max_rows) if truncated else df
    
    # 处理列名
    cols = [str(c).replace('\n', ' ').replace('|', '\\|') for c in df_show.columns]
    lines = [
        '| ' + ' | '.join(cols) + ' |',
        '|' + '|'.join(['---'] * len(cols)) + '|'
    ]
    
    # 使用 numpy 数组批量处理，比 iterrows 快 10~50 倍
    arr = df_show.astype(str).to_numpy()
    for row in arr:
        cells = [str(v).replace('\n', '<br>').replace('|', '\\|') for v in row]
        lines.append('| ' + ' | '.join(cells) + ' |')
        
    if truncated:
        lines.append(f"\n[已截断：仅显示前 {max_rows} 行，共 {n_rows} 行]")
        
    return '\n'.join(lines)


# ============ 格式处理器 ============
class FormatHandler:
    @classmethod
    def can_handle(cls, ctx: ReadContext) -> bool:
        raise NotImplementedError
    
    @classmethod
    def read(cls, ctx: ReadContext) -> FileReadResult:
        raise NotImplementedError


class TextHandler(FormatHandler):
    """纯文本处理器"""
    TEXT_EXTENSIONS = {
        '.txt', '.md', '.markdown', '.rst', '.py', '.js', '.ts', '.jsx', '.tsx', '.vue',
        '.json', '.yaml', '.yml', '.xml', '.html', '.htm', '.css', '.scss', '.less',
        '.sh', '.bash', '.zsh', '.fish', '.ps1', '.bat', '.cmd',
        '.sql', '.c', '.cpp', '.h', '.hpp', '.java', '.go', '.rs', '.rb', '.php',
        '.swift', '.kt', '.scala', '.r', '.m', '.mm', '.pl', '.lua', '.vim',
        '.toml', '.ini', '.cfg', '.conf', '.properties', '.env', '.gitignore',
        '.log', '.svg', '.graphql', '.proto', '.dart', '.zig', '.nim'
    }
    
    @classmethod
    def can_handle(cls, ctx: ReadContext) -> bool:
        if ctx.path.suffix.lower() in cls.TEXT_EXTENSIONS:
            return True
        if ctx.mime_type and ctx.mime_type.startswith(('text/', 'application/json', 'application/xml', 'application/javascript')):
            return True
        return False
    
    @classmethod
    def read(cls, ctx: ReadContext) -> FileReadResult:
        raw = ctx.path.read_bytes()
        candidates = []
        
        if ctx.encoding:
            candidates.append(ctx.encoding)
        candidates.extend(['utf-8-sig', 'utf-8'])
        
        chardet = _get_chardet()
        if chardet and len(raw) > 512:  # 太短 chardet 不准
            detected = chardet.detect(raw).get('encoding')
            if detected:
                candidates.append(detected)
        candidates.extend(['gbk', 'gb18030', 'latin-1'])
        
        actual_enc = 'utf-8 (with replacement)'
        content = None
        for enc in candidates:
            try:
                content = raw.decode(enc)
                actual_enc = enc
                break
            except (UnicodeDecodeError, LookupError):
                continue
                
        if content is None:
            content = raw.decode('utf-8', errors='replace')
        
        return FileReadResult(
            content=content,
            format='text',
            mime_type=ctx.mime_type or 'text/plain',
            metadata={'encoding': actual_enc, 'size_bytes': len(raw)}
        )


class CSVHandler(FormatHandler):
    """CSV/TSV 处理器"""
    @classmethod
    def can_handle(cls, ctx: ReadContext) -> bool:
        return ctx.path.suffix.lower() in {'.csv', '.tsv'}
    
    @classmethod
    def read(cls, ctx: ReadContext) -> FileReadResult:
        pd = _get_pandas()
        if pd is None:
            # 降级为文本处理
            return TextHandler.read(ctx)
        
        enc = ctx.encoding or 'utf-8-sig'
        sep = '\t' if ctx.path.suffix.lower() == '.tsv' else ','
        
        try:
            df = pd.read_csv(ctx.path, encoding=enc, sep=sep)
        except UnicodeDecodeError:
            for try_enc in ['gbk', 'gb18030', 'latin-1']:
                try:
                    df = pd.read_csv(ctx.path, encoding=try_enc, sep=sep)
                    enc = try_enc
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise FileReadError(f"无法解码 CSV 文件: {ctx.path}", code='DECODE_FAIL')
        
        return FileReadResult(
            content=df_to_markdown(df),
            format='markdown',
            mime_type='text/markdown',
            metadata={
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': list(df.columns),
                'encoding': enc,
            }
        )


class ExcelHandler(FormatHandler):
    """Excel 处理器，支持多 sheet，修复重复打开文件问题"""
    @classmethod
    def can_handle(cls, ctx: ReadContext) -> bool:
        return ctx.path.suffix.lower() in {'.xlsx', '.xls', '.xlsm', '.xlsb'}
    
    @classmethod
    def read(cls, ctx: ReadContext) -> FileReadResult:
        pd = _get_pandas()
        if pd is None:
            raise ImportError("读取 Excel 需要安装 pandas: pip install pandas openpyxl")
        
        ext = ctx.path.suffix.lower()
        engine = 'xlrd' if ext == '.xls' else 'openpyxl'
        
        try:
            # 复用 ExcelFile 对象，避免多次磁盘 I/O
            with pd.ExcelFile(ctx.path, engine=engine) as xl:
                sheets_info = xl.sheet_names
                
                if ctx.sheet_name is not None:
                    df = pd.read_excel(xl, sheet_name=ctx.sheet_name)
                    content = df_to_markdown(df)
                    meta_sheets = [ctx.sheet_name]
                    meta_rows = len(df)
                    meta_cols = len(df.columns)
                else:
                    if len(sheets_info) == 1:
                        df = pd.read_excel(xl, sheet_name=0)
                        content = df_to_markdown(df)
                    else:
                        parts = []
                        for name in sheets_info:
                            df_sheet = pd.read_excel(xl, sheet_name=name)
                            parts.append(f"## Sheet: {name}\n\n{df_to_markdown(df_sheet)}\n\n")
                        content = ''.join(parts)
                    
                    meta_sheets = sheets_info
                    meta_rows = len(content.splitlines())
                    meta_cols = 0
                    
            return FileReadResult(
                content=content,
                format='markdown',
                mime_type='text/markdown',
                metadata={
                    'sheets': meta_sheets,
                    'rows': meta_rows,
                    'columns': meta_cols,
                }
            )
                
        except ImportError:
            raise ImportError(f"读取 {ext} 需要安装对应引擎：pip install {engine}") from None
        except Exception as e:
            md = _get_markitdown()
            if md:
                try:
                    result = md.convert(str(ctx.path))
                    return FileReadResult(
                        content=result.text_content,
                        format='markdown',
                        mime_type='text/markdown',
                        metadata={'fallback': 'markitdown', 'sheets': []}
                    )
                except Exception:
                    pass
            raise FileReadError(f"Excel 读取失败: {e}", code='EXCEL_READ_FAIL') from e


class DocxHandler(FormatHandler):
    """Word 文档处理器，修复表头分隔行逻辑"""
    @classmethod
    def can_handle(cls, ctx: ReadContext) -> bool:
        return ctx.path.suffix.lower() == '.docx'

    @classmethod
    def read(cls, ctx: ReadContext) -> FileReadResult:
        docx = _get_docx()
        if docx is None:
            # 尝试 markitdown 兜底
            md = _get_markitdown()
            if md:
                try:
                    result = md.convert(str(ctx.path))
                    return FileReadResult(
                        content=result.text_content,
                        format='markdown',
                        mime_type='text/markdown',
                        metadata={'fallback': 'markitdown'}
                    )
                except Exception:
                    pass
            raise ImportError("读取 DOCX 需要安装 python-docx: pip install python-docx")

        try:
            doc = docx.Document(ctx.path)
        except Exception as e:
            # python-docx 失败时尝试 markitdown 兜底
            md = _get_markitdown()
            if md:
                try:
                    result = md.convert(str(ctx.path))
                    return FileReadResult(
                        content=result.text_content,
                        format='markdown',
                        mime_type='text/markdown',
                        metadata={'fallback': 'markitdown', 'error': str(e)}
                    )
                except Exception:
                    pass
            raise FileReadError(f"DOCX 读取失败: {e}", code='DOCX_READ_FAIL') from e

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = '\n\n'.join(paragraphs)

        tables_md = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                # 清理换行符和管道符，防止 Markdown 表格断裂
                cells = [cell.text.replace('\n', '<br>').replace('|', '\\|') for cell in row.cells]
                rows.append('| ' + ' | '.join(cells) + ' |')

            if rows:
                # 修复: 根据实际第一行的单元格数生成分隔行
                n_cols = rows[0].count('|') - 1
                if n_cols > 0:
                    header_sep = '|' + '|'.join(['---'] * n_cols) + '|'
                    rows.insert(1, header_sep)
                tables_md.append('\n'.join(rows))

        if tables_md:
            content += '\n\n' + '\n\n'.join(tables_md)

        return FileReadResult(
            content=content,
            format='markdown',
            mime_type='text/markdown',
            metadata={'paragraphs': len(paragraphs), 'tables': len(doc.tables)}
        )


class PDFHandler(FormatHandler):
    """PDF 处理器，增加页数限制"""
    MAX_PAGES = 200

    @classmethod
    def can_handle(cls, ctx: ReadContext) -> bool:
        return ctx.path.suffix.lower() == '.pdf'

    @classmethod
    def read(cls, ctx: ReadContext) -> FileReadResult:
        PyPDF2 = _get_pypdf()
        if PyPDF2 is None:
            # 尝试 markitdown 兜底
            md = _get_markitdown()
            if md:
                try:
                    result = md.convert(str(ctx.path))
                    return FileReadResult(
                        content=result.text_content,
                        format='markdown',
                        mime_type='text/markdown',
                        metadata={'fallback': 'markitdown'}
                    )
                except Exception:
                    pass
            raise ImportError("读取 PDF 需要安装 PyPDF2: pip install PyPDF2")

        text_parts = []
        total_pages = 0
        try:
            with open(ctx.path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                total_pages = len(reader.pages)
                pages_to_read = reader.pages[:cls.MAX_PAGES]

                for i, page in enumerate(pages_to_read):
                    try:
                        text = page.extract_text()
                        if text:
                            text_parts.append(f"## Page {i + 1}\n\n{text}")
                    except Exception:
                        text_parts.append(f"## Page {i + 1}\n\n[无法提取文本]")

            if total_pages > cls.MAX_PAGES:
                text_parts.append(f"\n[已截断：仅显示前 {cls.MAX_PAGES} 页，共 {total_pages} 页]")
        except Exception as e:
            # PyPDF2 失败时尝试 markitdown 兜底
            md = _get_markitdown()
            if md:
                try:
                    result = md.convert(str(ctx.path))
                    return FileReadResult(
                        content=result.text_content,
                        format='markdown',
                        mime_type='text/markdown',
                        metadata={'fallback': 'markitdown', 'error': str(e)}
                    )
                except Exception:
                    pass
            raise FileReadError(f"PDF 读取失败: {e}", code='PDF_READ_FAIL') from e

        return FileReadResult(
            content='\n\n'.join(text_parts),
            format='markdown',
            mime_type='text/markdown',
            metadata={'pages': total_pages, 'truncated': total_pages > cls.MAX_PAGES}
        )


class MarkdownHandler(FormatHandler):
    """通用文档转 Markdown 处理器，仅作为兜底处理 ppt/doc 等"""
    SUPPORTED_EXTS = {'.pptx', '.ppt', '.doc'}  # 移除已有专门处理器的扩展名
    
    @classmethod
    def can_handle(cls, ctx: ReadContext) -> bool:
        return ctx.path.suffix.lower() in cls.SUPPORTED_EXTS
    
    @classmethod
    def read(cls, ctx: ReadContext) -> FileReadResult:
        md = _get_markitdown()
        if md is None:
            raise ImportError(
                f"读取 {ctx.path.suffix} 文件需要安装 markitdown: pip install 'markitdown[all]'"
            )
        
        try:
            result = md.convert(str(ctx.path))
        except ImportError:
            raise
        except Exception as e:
            raise FileReadError(f"markitdown 转换失败: {e}", code='MARKITDOWN_FAIL') from e
        
        return FileReadResult(
            content=result.text_content,
            format='markdown',
            mime_type='text/markdown',
            metadata={'title': result.title or '', 'source': str(ctx.path)}
        )


class ImageHandler(FormatHandler):
    """图片处理器"""
    IMAGE_EXTS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp', '.tiff', '.ico'}
    
    @classmethod
    def can_handle(cls, ctx: ReadContext) -> bool:
        if ctx.path.suffix.lower() in cls.IMAGE_EXTS:
            return True
        if ctx.mime_type and ctx.mime_type.startswith('image/'):
            return True
        return False
    
    @classmethod
    def read(cls, ctx: ReadContext) -> FileReadResult:
        mime_type = ctx.mime_type or 'image/png'
        
        if ctx.return_base64:
            data = ctx.path.read_bytes()
            b64 = base64.b64encode(data).decode('ascii')
            content = f"data:{mime_type};base64,{b64}"
            size = len(data)
        else:
            # 返回绝对路径，让大模型知道去哪里找
            abs_path = ctx.path.resolve().as_posix()
            content = f"![{ctx.path.name}]({abs_path})\n\n(图片文件，路径: {abs_path})"
            size = ctx.path.stat().st_size
            
        return FileReadResult(
            content=content,
            format='image',
            mime_type=mime_type,
            metadata={'size_bytes': size}
        )


class BinaryHandler(FormatHandler):
    """二进制文件回退处理器"""
    @classmethod
    def can_handle(cls, ctx: ReadContext) -> bool:
        return True 
    
    @classmethod
    def read(cls, ctx: ReadContext) -> FileReadResult:
        mime_type = ctx.mime_type or 'application/octet-stream'
        content = (
            f"Binary file ({mime_type})\n"
            f"Size: {ctx.file_size} bytes\n"
            f"Path: {ctx.path}\n"
            f"[二进制内容，无法以文本形式呈现]"
        )
        
        return FileReadResult(
            content=content,
            format='binary',
            mime_type=mime_type,
            metadata={'size_bytes': ctx.file_size}
        )


# ============ 处理器优先级列表 (已修复) ============
HANDLERS = [
    ImageHandler,
    CSVHandler,
    ExcelHandler,
    DocxHandler,
    PDFHandler,
    TextHandler,
    MarkdownHandler,  # 兜底处理 ppt/doc 等
    BinaryHandler,    # 最终兜底
]


# ============ 并发控制 ============
_SEMAPHORE = asyncio.Semaphore(4)  # 限制同时读取 4 个文件，防止耗尽线程池


# ============ 主函数 ============
async def file_read(
    path: str,
    sheet_name: Optional[Union[str, int]] = None,
    encoding: str = "UTF-8",
    max_size_mb: int = 10,
    return_image_base64: bool = False,
    allowed_dirs: Optional[List[str]] = None
) -> Dict[str, Any]:
    """
    读取指定路径的文件内容
    """
    # 1. 路径与权限校验
    if allowed_dirs is None:
        allowed_dirs = [config.uploads_dir, backend.workspace_path]
    allowed_paths = [Path(p).resolve() for p in allowed_dirs]

    max_bytes = max_size_mb * 1024 * 1024
    path_obj = Path(path)
    
    if not path_obj.is_absolute():
        path_obj = Path.cwd() / path_obj
    
    try:
        safe_path = validate_path(str(path_obj), allowed_paths)
        if not isinstance(safe_path, Path):
            safe_path = Path(safe_path)
    except ValueError as e:
        raise FileReadError(f"路径校验失败：{e}", code='PATH_DENIED') from e

    if not safe_path.exists():
        raise FileReadError(f"文件不存在：{safe_path}", code='NOT_FOUND')
    if not safe_path.is_file():
        raise FileReadError(f"路径不是文件：{safe_path}", code='NOT_FILE')

    # 2. 文件大小检查
    try:
        file_size = safe_path.stat().st_size
    except OSError as e:
        raise FileReadError(f"无法获取文件信息：{e}", code='STAT_FAIL') from e

    if file_size > max_bytes:
        raise FileReadError(
            f"文件过大（{file_size} 字节），超过限制 {max_bytes} 字节。",
            code='FILE_TOO_LARGE'
        )

    # 3. MIME 类型猜测
    mime_type, _ = mimetypes.guess_type(str(safe_path))
    mime_type = mime_type or 'application/octet-stream'

    # 4. 构建上下文
    ctx = ReadContext(
        path=safe_path,
        mime_type=mime_type,
        encoding=encoding,
        sheet_name=sheet_name,
        return_base64=return_image_base64,
        file_size=file_size
    )

    loop = asyncio.get_running_loop()

    def _sync_read():
        last_error = None
        for handler in HANDLERS:
            if not handler.can_handle(ctx):
                continue
            try:
                result = handler.read(ctx)
                # 统一补充通用元信息
                result.metadata.setdefault('size_bytes', file_size)
                result.metadata.setdefault('path', str(safe_path))
                return asdict(result)
            except ImportError as e:
                last_error = e
                continue
            except FileReadError:
                raise
            except Exception as e:
                if handler is BinaryHandler:
                    raise FileReadError(f"读取文件时发生致命错误：{e}", code='READ_FAIL') from e
                last_error = e
                continue
                
        if last_error:
            raise FileReadError(
                f"缺少必要的依赖或处理失败，无法读取该文件: {last_error}",
                code='NO_HANDLER_OR_DEP_FAIL'
            ) from last_error
        raise FileReadError("无法读取文件：没有合适的处理器", code='NO_HANDLER')

    # 5. 放入线程池执行，并使用信号量控制并发
    async with _SEMAPHORE:
        return await loop.run_in_executor(None, _sync_read)
